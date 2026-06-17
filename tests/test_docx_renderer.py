"""Tests del exportador a Word (.docx).

Renderizamos un manual y volvemos a ABRIR el .docx con python-docx para verificar
que el Markdown se mapeó a estilos NATIVOS de Word (headings, listas, código,
tablas), no a texto plano. Es la verificación real: el archivo se puede abrir.
"""
from io import BytesIO

from docx import Document

from src.domain.entities import Manual, ManualType, ManualVersion, Section
from src.infrastructure.docx.python_docx_renderer import PythonDocxRenderer

_MD = """# Manual de Prueba

Este es un **párrafo** con `código` inline.

## Pasos a seguir

- Primero
- Segundo
  - Anidado

1. Uno
2. Dos

```vba
Sub Foo()
  MsgBox 1
End Sub
```

| Campo | Valor |
|-------|-------|
| Sitio | https://contoso.sharepoint.com |
| Lista | abc-123 |
"""


def _manual_y_version(md: str = _MD) -> tuple[Manual, ManualVersion]:
    manual = Manual(title="Cuentas por Pagar", type=ManualType.TECNICO, category="Power Automate")
    version = ManualVersion(
        version=2, content_html="<p>x</p>",
        sections=[Section(title="", content_html="<p>x</p>", order=0, source_markdown=md)],
    )
    return manual, version


def _render(md: str = _MD, **identity) -> Document:
    manual, version = _manual_y_version(md)
    data = PythonDocxRenderer(**identity).render(manual, version)
    assert data[:2] == b"PK"  # un .docx es un ZIP: empieza con PK
    return Document(BytesIO(data))


def test_genera_docx_valido_con_portada():
    doc = _render(brand="ACME Corp")
    full = "\n".join(p.text for p in doc.paragraphs)
    assert "ACME Corp" in full                  # marca en la portada
    assert "Cuentas por Pagar" in full          # título del manual
    assert "Manual Técnico" in full             # tipo
    assert "Power Automate" in full             # categoría


def test_headings_son_estilos_nativos():
    doc = _render()
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert any("Pasos a seguir" in h for h in headings)   # ## -> Heading 2 navegable


def test_listas_y_parrafo_con_formato():
    doc = _render()
    list_paras = [p.text for p in doc.paragraphs if "List" in p.style.name]
    assert "Primero" in list_paras and "Segundo" in list_paras
    assert "Uno" in list_paras and "Dos" in list_paras
    full = "\n".join(p.text for p in doc.paragraphs)
    assert "párrafo" in full and "código" in full         # inline desenvuelto (sin ** ni `)
    assert "**" not in full and "`" not in full


def test_bloque_de_codigo_se_preserva():
    doc = _render()
    full = "\n".join(p.text for p in doc.paragraphs)
    assert "MsgBox 1" in full
    assert "Sub Foo()" in full
    assert "```" not in full                              # la cerca no se imprime


def test_tabla_markdown_se_vuelve_tabla_word():
    doc = _render()
    assert len(doc.tables) == 1
    cells = [c.text for row in doc.tables[0].rows for c in row.cells]
    assert "Campo" in cells and "Valor" in cells          # encabezado
    assert "https://contoso.sharepoint.com" in cells      # fila de datos
    assert "abc-123" in cells


def test_sin_source_markdown_cae_al_html_sin_romper():
    manual = Manual(title="X", type=ManualType.FUNCIONAL, category="Otro")
    version = ManualVersion(
        version=1, content_html="",
        sections=[Section(title="", content_html="<p>Hola <b>mundo</b></p>", order=0,
                          source_markdown="")],
    )
    data = PythonDocxRenderer().render(manual, version)
    full = "\n".join(p.text for p in Document(BytesIO(data)).paragraphs)
    assert "Hola" in full and "mundo" in full             # texto del HTML, sin tags
    assert "<p>" not in full
